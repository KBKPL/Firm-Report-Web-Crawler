import os
import subprocess
import tempfile
import logging
import re
from typing import List, Dict, Optional, Tuple
from bs4 import BeautifulSoup
from docx import Document

from src.crawlers.base import CompanyCrawler
from src.utils.http_utils import session
from src.utils.pdf_utils import download_pdf
from src.utils.text_utils import sanitize_text, find_paragraphs_with_keyword, is_chinese_char, generate_variants
from src.utils.docx_utils import  add_keyword_variant_paragraphs

logger = logging.getLogger(__name__)

class GanfengCrawler(CompanyCrawler):
    """Crawler for Ganfeng Lithium quarterly reports and announcements."""
    SECTIONS = {
        "1": ("业绩报告", "quarterly performance", "crawl_quarterly_performance"),
        "2": ("公司公告", "company announcements", "crawl_company_announcements"),
    }

    def __init__(self, full_code: str, config: dict):
        super().__init__(full_code, config)
        # base_url from config
        self.base_url = config.get("base_url")
        self.financial_base_url = config.get("financial_base_url")
        self.announcement_base_url = config.get("company_announcement_url")

    def fetch_quarterly_performance_page(self, page: int = 1) -> List[Dict[str, str]]:
        """Fetch one page of quarterly performance metadata via HTML."""
        url = self.financial_base_url.format(page=page)
        resp = session.get(url)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "lxml")
        items = soup.find_all("div", class_="iryeji")
        if not items:
            return []
        # items are direct report blocks
        
        records: List[Dict[str, str]] = []
        for item in items:
            a = item.find("a", class_="iryejia")
            if not a or not a.get("href"):
                continue
            title = a.get_text(strip=True)
            href = a["href"]
            detail_url = href if href.startswith("http") else self.base_url + href
            resp2 = session.get(detail_url)
            resp2.raise_for_status()
            soup2 = BeautifulSoup(resp2.text, "lxml")
            # extract publication date from <div class="dettime">
            date_div = soup2.find("div", class_="dettime")
            if date_div:
                # text like '发布时间：2025-04-30'
                raw = date_div.get_text(strip=True)
                pub_date = raw.replace("发布时间：", "").strip()
            else:
                pub_date = None
            pdf_tag = soup2.find("a", href=lambda h: h and h.lower().endswith(".pdf"))
            if not pdf_tag:
                continue
            pdf_href = pdf_tag["href"]
            pdf_url = pdf_href if pdf_href.startswith("http") else self.base_url + pdf_href
            records.append({"publishDate": pub_date, "title": title, "pdf_url": pdf_url})
        return records

    def crawl_quarterly_performance(
        self,
        keywords: List[str],
        output_dir: str,
        start_date: Optional[str] = None,
        end_date: Optional[str] = None,
    ) -> Dict[str, str]:
        os.makedirs(output_dir, exist_ok=True)
        docs = {kw: Document() for kw in keywords}
        generated: Dict[str, str] = {}
        page = 1
        while True:
            logger.info(f"Fetching page: {page}")
            records = self.fetch_quarterly_performance_page(page)
            if not records:
                break
            break_page = False
            for rec in records:
                pub_date = rec.get("publishDate")
                if end_date and pub_date and pub_date > end_date:
                    continue
                if start_date and pub_date and pub_date < start_date:
                    break_page = True
                    break
                title = rec.get("title", "")
                pdf_url = rec.get("pdf_url")
                pdf_bytes = download_pdf(pdf_url)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                    tmp_pdf.write(pdf_bytes)
                    tmp_pdf_path = tmp_pdf.name
                tmp_txt = tempfile.NamedTemporaryFile(delete=False, suffix=".txt")
                tmp_txt.close()
                try:
                    subprocess.run(["pdftext", tmp_pdf_path, "--out_path", tmp_txt.name], check=True)
                    with open(tmp_txt.name, "r", encoding="utf-8") as f_txt:
                        text = f_txt.read()
                except Exception as e:
                    logger.error(f"Error extracting text: {e}")
                    continue
                finally:
                    os.remove(tmp_pdf_path)
                    os.remove(tmp_txt.name)
                text = sanitize_text(text)
                for kw in keywords:
                    variants = generate_variants(kw)
                    paras: List[str] = []
                    for v in variants:
                        paras.extend(find_paragraphs_with_keyword(text, v))
                    paras = list(dict.fromkeys(paras))
                    paras = [p for p in paras if any(is_chinese_char(ch) for ch in p)]
                    if not paras:
                        continue
                    doc = docs[kw]
                    doc.add_heading(title, level=1)
                    add_keyword_variant_paragraphs(doc, paras, variants, pdf_url)
            if break_page:
                break
            page += 1
        for kw, doc in docs.items():
            safe_kw = kw.replace(" ", "_")
            out_name = f"{self.full_code}_{safe_kw}_quarterly_performance.docx"
            path = os.path.join(output_dir, out_name)
            doc.save(path)
            generated[kw] = path
        return generated

    def fetch_company_announcement_page(self, page: int = 1) -> List[Dict[str, str]]:
        """Fetch one page of company announcements metadata via HTML."""
        url = self.announcement_base_url.format(page=page)
        resp = session.get(url)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "lxml")
        items = soup.find_all("div", class_="irggfl")
        records: List[Dict[str, str]] = []
        for item in items:
            font_link = item.find("a", class_="irfont")
            if not font_link or not font_link.get("href"):
                continue
            title = font_link.get_text(strip=True)
            href = font_link["href"]
            detail_url = href if href.startswith("http") else self.base_url + href
            resp2 = session.get(detail_url)
            resp2.raise_for_status()
            soup2 = BeautifulSoup(resp2.text, "lxml")
            date_div = soup2.find("div", class_="dettime")
            if date_div:
                raw = date_div.get_text(strip=True)
                pub_date = raw.replace("发布时间：", "").strip()
            else:
                pub_date = None
            pdf_tag = soup2.find("a", href=lambda h: h and h.lower().endswith(".pdf"))
            if not pdf_tag:
                continue
            pdf_href = pdf_tag["href"]
            pdf_url = pdf_href if pdf_href.startswith("http") else self.base_url + pdf_href
            records.append({"publishDate": pub_date, "title": title, "pdf_url": pdf_url})
        return records

    def crawl_company_announcements(
        self,
        keywords: List[str],
        output_dir: str,
        start_date: Optional[str] = None,
        end_date: Optional[str] = None,
    ) -> Dict[str, str]:
        os.makedirs(output_dir, exist_ok=True)
        docs = {kw: Document() for kw in keywords}
        generated: Dict[str, str] = {}
        page = 1
        while True:
            logger.info(f"Fetching announcements page: {page}")
            records = self.fetch_company_announcement_page(page)
            if not records:
                break
            break_page = False
            for rec in records:
                pub_date = rec.get("publishDate")
                if end_date and pub_date and pub_date > end_date:
                    continue
                if start_date and pub_date and pub_date < start_date:
                    break_page = True
                    break
                title = rec.get("title", "")
                pdf_url = rec.get("pdf_url")
                pdf_bytes = download_pdf(pdf_url)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                    tmp_pdf.write(pdf_bytes)
                    tmp_pdf_path = tmp_pdf.name
                tmp_txt = tempfile.NamedTemporaryFile(delete=False, suffix=".txt")
                tmp_txt.close()
                try:
                    subprocess.run(["pdftext", tmp_pdf_path, "--out_path", tmp_txt.name], check=True)
                    with open(tmp_txt.name, "r", encoding="utf-8") as f_txt:
                        text = f_txt.read()
                except Exception as e:
                    logger.error(f"Error extracting text: {e}")
                    continue
                finally:
                    os.remove(tmp_pdf_path)
                    os.remove(tmp_txt.name)
                text = sanitize_text(text)
                for kw in keywords:
                    variants = generate_variants(kw)
                    paras: List[str] = []
                    for v in variants:
                        paras.extend(find_paragraphs_with_keyword(text, v))
                    paras = list(dict.fromkeys(paras))
                    paras = [p for p in paras if any(is_chinese_char(ch) for ch in p)]
                    if not paras:
                        continue
                    doc = docs[kw]
                    doc.add_heading(title, level=1)
                    add_keyword_variant_paragraphs(doc, paras, variants, pdf_url)
            if break_page:
                break
            page += 1
        for kw, doc in docs.items():
            safe_kw = kw.replace(" ", "_")
            out_name = f"{self.full_code}_{safe_kw}_company_announcements.docx"
            path = os.path.join(output_dir, out_name)
            doc.save(path)
            generated[kw] = path
        return generated
