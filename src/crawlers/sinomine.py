import os
import logging
import base64
from typing import List, Dict, Optional
from docx import Document
from bs4 import BeautifulSoup

from src.crawlers.base import CompanyCrawler
from src.utils.docx_utils import add_keyword_paragraphs
from src.utils.http_utils import session
from src.utils.pdf_utils import download_pdf
from src.utils.text_utils import find_paragraphs_with_keyword, sanitize_text

logger = logging.getLogger(__name__)

class SinomineCrawler(CompanyCrawler):
    SECTIONS = {
        "1": ("季度业绩", "quarterly performance", "crawl_quarterly_performance"),
        "2": ("公司公告", "company announcements", "crawl_company_announcements"),
        "3": ("券商报告", "broker reports", "crawl_broker_reports"),
    }

    def __init__(self, full_code: str, config: dict):
        super().__init__(full_code, config)
        self.page_size = config.get("page_size", 10)
        self.broker_report_url = config.get("broker_report_url")
        self.broker_report_nonpdf_base_url = config.get("broker_report_nonpdf_base_url")
        self.company_announcement = config.get("company_announcement_url")
        self.financial_base_url = config.get("financial_base_url")
        self.store_id = config.get("store_id")

    def crawl_broker_reports(
        self,
        keywords: List[str],
        output_dir: str,
        start_date: Optional[str] = None,
        end_date: Optional[str] = None,
    ) -> Dict[str, str]:
        """Crawl broker reports and return mapping of keyword to output file path."""
        import subprocess
        import tempfile
        from src.utils.html_utils import fetch_rendered_html

        os.makedirs(output_dir, exist_ok=True)
        docs = {kw: Document() for kw in keywords}
        generated = {}
        page_index = 0
        while True:
            logger.info(f"Fetching page: {page_index}")
            records = self.fetch_broker_report_page(page_index)
            if not records:
                break
            break_page = False
            for rec in records:
                pub_date = rec.get("publishDate", "").split()[0]
                if end_date and pub_date > end_date:
                    continue
                if start_date and pub_date and pub_date < start_date:
                    break_page = True
                    break
                rec_id = rec.get("id")
                raw_url = rec.get("url")
                if not raw_url or '/report/detail' in raw_url:
                    detail_url = raw_url or f"{self.broker_report_nonpdf_base_url}?id={rec.get('reportId') or rec_id}&type={rec.get('type')}&storeId={self.store_id}"
                    html_text = fetch_rendered_html(detail_url)
                    if not html_text:
                        continue
                    text = BeautifulSoup(html_text, "lxml").get_text("\n")
                    url_used = detail_url
                else:
                    if 'onlinePreview' in raw_url:
                        preview_url = raw_url
                    else:
                        b64 = base64.urlsafe_b64encode(raw_url.encode()).decode()
                        preview_url = (
                            f"https://file-view.comein.cn/onlinePreview?url={b64}"
                            "&officePreviewSwitchDisabled=true"
                            "&officePreviewType=pdf"
                            "&watermarkTxt="
                        )
                    pdf_bytes = download_pdf(preview_url)
                    is_html = pdf_bytes.lstrip().startswith(b'<')
                    if is_html:
                        text = pdf_bytes.decode('utf-8', errors='ignore')
                    else:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                            tmp_pdf.write(pdf_bytes)
                            tmp_pdf_path = tmp_pdf.name
                        tmp_txt = tempfile.NamedTemporaryFile(delete=False, suffix=".txt")
                        tmp_txt.close()
                        try:
                            subprocess.run(["pdftext", tmp_pdf_path, "--out_path", tmp_txt.name], check=True)
                            with open(tmp_txt.name, "r", encoding="utf-8") as f_txt:
                                text = f_txt.read()
                        except Exception:
                            os.remove(tmp_pdf_path)
                            os.remove(tmp_txt.name)
                            continue
                        finally:
                            os.remove(tmp_pdf_path)
                            os.remove(tmp_txt.name)
                    url_used = preview_url
                for kw in keywords:
                    text = sanitize_text(text)
                    paras = find_paragraphs_with_keyword(text, kw)
                    if not paras:
                        continue
                    doc = docs[kw]
                    title = rec.get("title", "")
                    author = rec.get("author", "")
                    doc.add_heading(f"{pub_date}_{title}_{author}", level=1)
                    add_keyword_paragraphs(doc, paras, kw, url_used)
            if len(records) < self.page_size or break_page:
                break
            page_index += 1
        # save docs for broker reports
        for kw, doc in docs.items():
            safe_kw = kw.replace(" ", "_")
            out_name = f"{self.full_code}_{safe_kw}_券商报告.docx"
            path = os.path.join(output_dir, out_name)
            doc.save(path)
            logger.info(f"Saved combined DOCX for keyword {kw} [券商报告]: {out_name}")
            generated[kw] = path
        return generated

    def crawl_company_announcements(
        self,
        keywords: List[str],
        output_dir: str,
        start_date: Optional[str] = None,
        end_date: Optional[str] = None,
    ) -> Dict[str, str]:
        """Crawl 公司公告 section and extract keyword-containing paragraphs."""
        import os, tempfile, subprocess
        from docx import Document
        # download_pdf and keyword finder already imported at top

        os.makedirs(output_dir, exist_ok=True)
        docs = {kw: Document() for kw in keywords}
        generated = {}
        page_index = 0
        while True:
            logger.info(f"Fetching company report page: {page_index}")
            records = self.fetch_company_announcement_page(page_index)
            if not records:
                break
            break_page = False
            for rec in records:
                pub_date = rec.get("publishDate", "").split()[0]
                if end_date and pub_date > end_date:
                    continue
                if start_date and pub_date and pub_date < start_date:
                    break_page = True
                    break
                rec_id = rec.get('reportId') or rec.get('id')
                raw_url = rec.get('comeinLink') or rec.get('url')
                # if not raw_url or '/report/detail' in raw_url:
                #     detail_url = raw_url or f"{self.detail_base_url}?id={rec_id}&type={rec.get('type')}&storeId={self.store_id}"
                #     html_text = fetch_rendered_html(detail_url)
                #     if not html_text:
                #         continue
                #     text = BeautifulSoup(html_text, "lxml").get_text("\n")
                #     url_used = detail_url
                # else:
                #     if 'onlinePreview' in raw_url:
                preview_url = raw_url
                pdf_bytes = download_pdf(preview_url)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                    tmp_pdf.write(pdf_bytes)
                    tmp_pdf_path = tmp_pdf.name
                tmp_txt = tempfile.NamedTemporaryFile(delete=False, suffix=".txt")
                tmp_txt.close()
                try:
                    subprocess.run(["pdftext", tmp_pdf_path, "--out_path", tmp_txt.name], check=True)
                    with open(tmp_txt.name, "r", encoding="utf-8") as f_txt:
                        text = f_txt.read()
                except Exception:
                    os.remove(tmp_pdf_path)
                    os.remove(tmp_txt.name)
                    continue
                finally:
                    os.remove(tmp_pdf_path)
                    os.remove(tmp_txt.name)
                url_used = preview_url
                for kw in keywords:
                    text = sanitize_text(text)
                    paras = find_paragraphs_with_keyword(text, kw)
                    if not paras:
                        continue
                    doc = docs[kw]
                    title = rec.get("title", "")
                    author = rec.get("author", "")
                    doc.add_heading(f"{pub_date}_{title}_{author}", level=1)
                    add_keyword_paragraphs(doc, paras, kw, url_used)
            if len(records) < self.page_size or break_page:
                break
            page_index += 1
        # save docs
        for kw, doc in docs.items():
            safe_kw = kw.replace(" ", "_")
            out_name = f"{self.full_code}_{safe_kw}_公司公告.docx"
            doc.save(os.path.join(output_dir, out_name))
            logger.info(f"Saved combined DOCX for keyword {kw} [公司公告]: {out_name}")
            generated[kw] = os.path.join(output_dir, out_name)
        return generated

    def crawl_quarterly_performance(
        self,
        keywords: List[str],
        output_dir: str,
        start_date: Optional[str] = None,
        end_date: Optional[str] = None,
    ) -> Dict[str, str]:
        """Crawl 季度业绩 using Playwright UI, extract keyword paragraphs."""
        import os, tempfile, subprocess
        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            logger.error("Playwright not installed. Run 'pip install playwright' and 'playwright install'.")
            return {}
        os.makedirs(output_dir, exist_ok=True)
        docs = {kw: Document() for kw in keywords}
        generated = {}
        # fetch JSON selections
        records = self.fetch_financial_statement_page("", "")
        selections: list[tuple[str, str]] = []
        for rec in records:
            pub = rec.get("publishDate", "").split()[0]
            if not pub:
                continue
            if end_date and pub > end_date:
                continue
            if start_date and pub < start_date:
                continue
            year = pub.split("-")[0]
            year_int = int(year)
            md = pub[5:]
            if md == '12-31':
                rpt_cn = '年度'
            elif year_int >= 2015 and md == '09-30':
                rpt_cn = '三季度'
            elif year_int >= 2015 and md == '06-30':
                rpt_cn = '半年度'
            elif year_int >= 2015 and md == '03-31':
                rpt_cn = '一季度'
            else:
                continue
            selections.append((year, rpt_cn))
        if not selections:
            logger.info("No quarterly reports found in date range.")
            return {}
        # drive UI
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            url = (
                "https://irm-enterprise-pc.comein.cn/investors/flow/mme/15707/"
                "quarterly_performance?lang=zh-cn&storeId=21113"
            )
            page.goto(url, timeout=60000)
            page.wait_for_load_state("networkidle")
            for year, rpt_cn in selections:
                logger.info(f"Selecting {year} {rpt_cn}")
                while page.locator(f"text={year}").count() == 0:
                    page.click(".el-carousel__arrow--right")
                    page.wait_for_timeout(500)
                page.click(f"text={year}")
                page.click(f"text={rpt_cn}")
                with page.expect_popup() as pop_info:
                    page.click("text=业绩公告")
                pdf_page = pop_info.value
                pdf_url = pdf_page.url
                pdf_page.close()
                # download and extract
                pdf_bytes = download_pdf(pdf_url)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                    tmp_pdf.write(pdf_bytes)
                    tmp_pdf_path = tmp_pdf.name
                tmp_txt = tempfile.NamedTemporaryFile(delete=False, suffix=".txt")
                tmp_txt.close()
                try:
                    subprocess.run(["pdftext", tmp_pdf_path, "--out_path", tmp_txt.name], check=True)
                    with open(tmp_txt.name, "r", encoding="utf-8") as f_txt:
                        text = f_txt.read()
                except Exception as e:
                    os.remove(tmp_pdf_path)
                    os.remove(tmp_txt.name)
                    logger.error(f"Failed extract text from PDF: {e}")
                    continue
                finally:
                    os.remove(tmp_pdf_path)
                    os.remove(tmp_txt.name)
                # scan keywords
                for kw in keywords:
                    text = sanitize_text(text)
                    paras = find_paragraphs_with_keyword(text, kw)
                    if not paras:
                        continue
                    doc = docs[kw]
                    doc.add_heading(f"{self.full_code}_{kw}_{year}_{rpt_cn}_{pub}", level=1)
                    add_keyword_paragraphs(doc, paras, kw, pdf_url)
            browser.close()
        # save docs
        for kw, doc in docs.items():
            safe = kw.replace(' ', '_')
            out = f"{self.full_code}_{safe}_季度业绩.docx"
            doc.save(os.path.join(output_dir, out))
            logger.info(f"Saved combined DOCX for keyword {kw} [季度业绩]: {out}")
            generated[kw] = os.path.join(output_dir, out)
        return generated

    def fetch_broker_report_page(self, page_index: int = 0) -> list[dict]:
        """Fetch one page of report metadata using 0-based page index (pagestart)."""
        payload = {
            "pagestart": page_index,  # page number (0-based)
            "pagenum": self.page_size,
            "fullCode": self.full_code,
            "keyword": "",
            "languageType": 0
        }
        try:
            resp = session.post(self.broker_report_url, json=payload, timeout=10)
            resp.raise_for_status()
            data = resp.json()
            if data.get("code") != "0":
                logging.error(f"List API error: {data}")
                return []
            return data.get("rows", [])
        except Exception as e:
            logging.error(f"Error fetching report list: {e}")
            return []

    def fetch_company_announcement_page(self, page_index: int = 0) -> list[dict]:
        """Fetch one page of 公司报告 metadata via GET."""
        url = self.company_announcement.format(full_code=self.full_code)
        params = {
            "classificationIds": "",
            "pageStart": page_index,
            "pageNum": self.page_size,
            "order": "desc",
            "title": "",
            "languageType": 0
        }
        try:
            resp = session.get(url, params=params, timeout=10)
            resp.raise_for_status()
            data = resp.json()
            if data.get("code") != "0":
                logging.error(f"Company report list API error: {data}")
                return []
            return data.get("rows", [])
        except Exception as e:
            logging.error(f"Error fetching company report page: {e}")
            return []

    def fetch_financial_statement_page(self, publish_date: str, report_type: str) -> list[dict]:
        """Fetch financial statement data for a given year and report type."""
        url = self.financial_base_url.format(full_code=self.full_code)
        params = {"publishDate": publish_date, "reportType": report_type}
        try:
            resp = session.get(url, params=params, timeout=10)
            resp.raise_for_status()
            data = resp.json()
            if data.get("code") != "0":
                logging.error(f"Financial statement API error: {data}")
                return []
            rows = data.get("rows") or data.get("data")
            if not rows:
                return []
            return rows if isinstance(rows, list) else [rows]
        except Exception as e:
            logging.error(f"Error fetching financial statement: {e}")
            return []
