"""
Crawler for Sinomine (company code sz002738) reports.
Fetches report metadata via JSON API and downloads each PDF for keyword extraction.
"""
import os
import logging
import subprocess
import base64
import re
import sys

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s: %(message)s")

from http_utils import session
from pdf_utils import download_pdf
from text_utils import find_paragraphs_with_keyword
from html_utils import find_paragraphs_from_html
from docx_utils import write_paragraphs_to_docx
from bs4 import BeautifulSoup  # for extracting full page text
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

API_LIST_URL = "https://server.comein.cn/comein/irmcenter/anonymous/irstore/report/list"
PAGE_SIZE = 10

# Base URL for HTML detail pages and default store ID for Sinomine
DETAIL_BASE_URL = "https://irm-enterprise-pc.comein.cn/investors/flow/report/detail"
STORE_ID = "21113"

# Endpoint for 公司报告
ANNOUNCE_URL = "https://server.comein.cn/comein/irmcenter/v3/anonymous/irstore/{full_code}/announcements"

# Base URL for 季度业绩 (quarter performance)
FINANCIAL_BASE_URL = "https://server.comein.cn/comein/irmcenter/anonymous/irstore/{full_code}/stock/financial-statement"

def fetch_rendered_html(url: str) -> str:
    """Use Playwright to render JS-driven pages and return full HTML."""
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        logging.error("Playwright not installed. Run 'pip install playwright' and 'playwright install'.")
        return None
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(url, timeout=60000)
            page.wait_for_load_state("networkidle")
            html = page.content()
            browser.close()
        return html
    except Exception as e:
        logging.error(f"Error rendering HTML via Playwright: {e}")
        return None

def save_page_as_pdf(url: str, output_path: str) -> bool:
    """Render a page and save it as PDF using Playwright."""
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        logging.error("Playwright not installed. Run 'pip install playwright' and 'playwright install'.")
        return False
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(url, timeout=60000, wait_until="networkidle")
            page.pdf(path=output_path, format="A4", print_background=True)
            browser.close()
        return True
    except Exception as e:
        logging.error(f"Error saving PDF via Playwright: {e}")
        return False

def fetch_report_page(full_code: str, page_index: int = 0, page_num: int = PAGE_SIZE) -> list[dict]:
    """Fetch one page of report metadata using 0-based page index (pagestart)."""
    payload = {
        "pagestart": page_index,  # page number (0-based)
        "pagenum": page_num,
        "fullCode": full_code,
        "keyword": "",
        "languageType": 0
    }
    try:
        resp = session.post(API_LIST_URL, json=payload, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        if data.get("code") != "0":
            logging.error(f"List API error: {data}")
            return []
        return data.get("rows", [])
    except Exception as e:
        logging.error(f"Error fetching report list: {e}")
        return []

def fetch_company_report_page(full_code: str, page_index: int = 0, page_num: int = PAGE_SIZE) -> list[dict]:
    """Fetch one page of 公司报告 metadata via GET."""
    url = ANNOUNCE_URL.format(full_code=full_code)
    params = {
        "classificationIds": "",
        "pageStart": page_index,
        "pageNum": page_num,
        "order": "desc",
        "title": "",
        "languageType": 0
    }
    try:
        resp = session.get(url, params=params, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        if data.get("code") != "0":
            logging.error(f"Company report list API error: {data}")
            return []
        return data.get("rows", [])
    except Exception as e:
        logging.error(f"Error fetching company report page: {e}")
        return []

def fetch_financial_statement_page(full_code: str, publish_date: str, report_type: str) -> list[dict]:
    """Fetch financial statement data for a given year and report type."""
    url = FINANCIAL_BASE_URL.format(full_code=full_code)
    params = {"publishDate": publish_date, "reportType": report_type}
    try:
        resp = session.get(url, params=params, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        if data.get("code") != "0":
            logging.error(f"Financial statement API error: {data}")
            return []
        rows = data.get("rows") or data.get("data")
        if not rows:
            return []
        return rows if isinstance(rows, list) else [rows]
    except Exception as e:
        logging.error(f"Error fetching financial statement: {e}")
        return []

def add_hyperlink(paragraph, url):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    i = OxmlElement('w:i')
    rPr.append(i)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = url
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def crawl_company(full_code: str, keywords: list[str], output_dir: str = "results", start_date: str = None, end_date: str = None):
    os.makedirs(output_dir, exist_ok=True)
    # perform a single crawl and accumulate docs per keyword
    docs = {kw: Document() for kw in keywords}
    page_index = 0
    while True:
        logging.info(f"Fetching page: {page_index}")
        records = fetch_report_page(full_code, page_index, PAGE_SIZE)
        if not records:
            break
        break_page = False
        for rec in records:
            pub_date = rec.get("publishDate", "").split()[0]
            if end_date and pub_date > end_date:
                continue
            if start_date and pub_date and pub_date < start_date:
                break_page = True
                break
            rec_id = rec.get('id')
            raw_url = rec.get('url')
            if not raw_url or '/report/detail' in raw_url:
                detail_url = raw_url or f"{DETAIL_BASE_URL}?id={rec.get('reportId') or rec_id}&type={rec.get('type')}&storeId={STORE_ID}"
                html_text = fetch_rendered_html(detail_url)
                if not html_text:
                    continue
                text = BeautifulSoup(html_text, "lxml").get_text("\n")
                url_used = detail_url
            else:
                if 'onlinePreview' in raw_url:
                    preview_url = raw_url
                else:
                    b64 = base64.urlsafe_b64encode(raw_url.encode()).decode()
                    preview_url = (
                        f"https://file-view.comein.cn/onlinePreview?url={b64}"
                        "&officePreviewSwitchDisabled=true"
                        "&officePreviewType=pdf"
                        "&watermarkTxt="
                    )
                pdf_bytes = download_pdf(preview_url)
                is_html = pdf_bytes.lstrip().startswith(b'<')
                if is_html:
                    text = pdf_bytes.decode('utf-8', errors='ignore')
                else:
                    with open("temp.pdf", "wb") as f_pdf:
                        f_pdf.write(pdf_bytes)
                    try:
                        subprocess.run(["pdftext", "temp.pdf", "--out_path", "temp.txt"], check=True)
                        with open("temp.txt", "r", encoding="utf-8") as f_txt:
                            text = f_txt.read()
                    except Exception:
                        continue
                url_used = preview_url
            # scan for each keyword
            for kw in keywords:
                paras = find_paragraphs_with_keyword(text, kw)
                if not paras:
                    continue
                doc = docs[kw]
                title = rec.get("title", "")
                author = rec.get("author", "")
                doc.add_heading(f"{pub_date}_{title}_{author}", level=1)
                p = doc.add_paragraph()
                add_hyperlink(p, url_used)
                for idx, para in enumerate(paras, start=1):
                    doc.add_heading(f"Location {idx}", level=2)
                    p2 = doc.add_paragraph()
                    pattern = re.compile(re.escape(kw), re.IGNORECASE)
                    pos = 0
                    for m in pattern.finditer(para):
                        if m.start() > pos:
                            p2.add_run(para[pos:m.start()])
                        run_h = p2.add_run(para[m.start():m.end()])
                        run_h.bold = True
                        run_h.font.highlight_color = WD_COLOR_INDEX.RED
                        pos = m.end()
                    if pos < len(para):
                        p2.add_run(para[pos:])
        if len(records) < PAGE_SIZE or break_page:
            break
        page_index += 1
    # save all combined docs
    for kw, doc in docs.items():
        safe_kw = kw.replace(" ", "_")
        out_name = f"{full_code}_{safe_kw}_券商研报.docx"
        doc.save(os.path.join(output_dir, out_name))
        logging.info(f"Saved combined DOCX for keyword {kw}: {out_name}")

def crawl_company_reports(full_code: str, keywords: list[str], output_dir: str = "results/company_reports", start_date: str = None, end_date: str = None):
    """Crawl 公司公告 section and extract keyword-containing paragraphs."""
    os.makedirs(output_dir, exist_ok=True)
    docs = {kw: Document() for kw in keywords}
    page_index = 0
    while True:
        logging.info(f"Fetching company report page: {page_index}")
        records = fetch_company_report_page(full_code, page_index, PAGE_SIZE)
        if not records:
            break
        break_page = False
        for rec in records:
            pub_date = rec.get("publishDate", "").split()[0]
            if end_date and pub_date > end_date:
                continue
            if start_date and pub_date and pub_date < start_date:
                break_page = True
                break
            rec_id = rec.get('reportId') or rec.get('id')
            # use comeinLink for announcements
            raw_url = rec.get('comeinLink') or rec.get('url')
            # HTML detail page when no preview blob
            if not raw_url or '/report/detail' in raw_url:
                detail_url = raw_url or f"{DETAIL_BASE_URL}?id={rec_id}&type={rec.get('type')}&storeId={STORE_ID}"
                html_text = fetch_rendered_html(detail_url)
                if not html_text:
                    continue
                text = BeautifulSoup(html_text, "lxml").get_text("\n")
                url_used = detail_url
            # otherwise direct PDF preview
            else:
                if 'onlinePreview' in raw_url:
                    preview_url = raw_url
                else:
                    b64 = base64.urlsafe_b64encode(raw_url.encode()).decode()
                    preview_url = (
                        f"https://file-view.comein.cn/onlinePreview?url={b64}"
                        "&officePreviewSwitchDisabled=true"
                        "&officePreviewType=pdf"
                        "&watermarkTxt="
                    )
                # download and extract text from PDF
                pdf_bytes = download_pdf(preview_url)
                with open("temp.pdf", "wb") as f_pdf:
                    f_pdf.write(pdf_bytes)
                try:
                    subprocess.run(["pdftext", "temp.pdf", "--out_path", "temp.txt"], check=True)
                    with open("temp.txt", "r", encoding="utf-8") as f_txt:
                        text = f_txt.read()
                except Exception:
                    continue
                url_used = preview_url
            for kw in keywords:
                paras = find_paragraphs_with_keyword(text, kw)
                if not paras:
                    continue
                doc = docs[kw]
                title = rec.get("title", "")
                author = rec.get("author", "")
                doc.add_heading(f"{pub_date}_{title}_{author}", level=1)
                p = doc.add_paragraph()
                add_hyperlink(p, url_used)
                for idx, para in enumerate(paras, start=1):
                    doc.add_heading(f"Location {idx}", level=2)
                    p2 = doc.add_paragraph()
                    pattern = re.compile(re.escape(kw), re.IGNORECASE)
                    pos = 0
                    for m in pattern.finditer(para):
                        if m.start() > pos:
                            p2.add_run(para[pos:m.start()])
                        run_h = p2.add_run(para[m.start():m.end()])
                        run_h.bold = True
                        run_h.font.highlight_color = WD_COLOR_INDEX.RED
                        pos = m.end()
                    if pos < len(para):
                        p2.add_run(para[pos:])
        if len(records) < PAGE_SIZE or break_page:
            break
        page_index += 1
    # save all combined docs
    for kw, doc in docs.items():
        safe_kw = kw.replace(' ', '_')
        out_name = f"{full_code}_{safe_kw}_公司公告.docx"
        doc.save(os.path.join(output_dir, out_name))
        logging.info(f"Saved combined DOCX for keyword {kw} [公司公告]: {out_name}")

def crawl_quarterly_reports(full_code: str, keywords: list[str], output_dir: str = "results/quarterly", start_date: str = None, end_date: str = None):
    """Crawl 季度业绩 using Playwright UI, extract keyword paragraphs."""
    os.makedirs(output_dir, exist_ok=True)
    docs = {kw: Document() for kw in keywords}
    # fetch JSON to build (year, quarter) selections
    records = fetch_financial_statement_page(full_code, "", "")
    selections: list[tuple[str, str]] = []
    for rec in records:
        pub = rec.get("publishDate", "").split()[0]
        if not pub:
            continue
        if end_date and pub > end_date:
            continue
        if start_date and pub < start_date:
            continue
        year = pub.split("-")[0]
        year_int = int(year)
        md = pub[5:]
        if md == '12-31':
            rpt_cn = '年度'
        elif year_int >= 2015 and md == '09-30':
            rpt_cn = '三季度'
        elif year_int >= 2015 and md == '06-30':
            rpt_cn = '半年度'
        elif year_int >= 2015 and md == '03-31':
            rpt_cn = '一季度'
        else:
            continue
        selections.append((year, rpt_cn))
    if not selections:
        logging.info("No quarterly reports found in date range.")
        return
    # use Playwright to drive UI and capture preview URLs
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        logging.error("Playwright not installed. Run 'pip install playwright' and 'playwright install'.")
        return
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        url = (
            "https://irm-enterprise-pc.comein.cn/investors/flow/mme/15707/"
            "quarterly_performance?lang=zh-cn&storeId=21113"
        )
        page.goto(url, timeout=60000)
        page.wait_for_load_state("networkidle")
        for year, rpt_cn in selections:
            logging.info(f"Selecting {year} {rpt_cn}")
            # reveal year if hidden
            while page.locator(f"text={year}").count() == 0:
                page.click(".el-carousel__arrow--right")
                page.wait_for_timeout(500)
            page.click(f"text={year}")
            page.click(f"text={rpt_cn}")
            # open PDF in popup
            with page.expect_popup() as pop_info:
                page.click("text=业绩公告")
            pdf_page = pop_info.value
            pdf_url = pdf_page.url
            pdf_page.close()
            # download & extract
            try:
                pdf_bytes = download_pdf(pdf_url)
                with open("temp.pdf", "wb") as f: f.write(pdf_bytes)
                subprocess.run(["pdftext", "temp.pdf", "--out_path", "temp.txt"], check=True)
                with open("temp.txt", "r", encoding="utf-8") as f: text = f.read()
            except Exception as e:
                logging.error(f"Failed PDF fetch/parse for {year} {rpt_cn}: {e}")
                continue
            # scan keywords
            for kw in keywords:
                paras = find_paragraphs_with_keyword(text, kw)
                if not paras: continue
                doc = docs[kw]
                doc.add_heading(f"{full_code}_{kw}_{year}_{rpt_cn}_{pub}", level=1)
                p_link = doc.add_paragraph(); add_hyperlink(p_link, pdf_url)
                for idx, para in enumerate(paras, start=1):
                    doc.add_heading(f"Location {idx}", level=2)
                    p2 = doc.add_paragraph()
                    pattern = re.compile(re.escape(kw), re.IGNORECASE)
                    pos = 0
                    for m in pattern.finditer(para):
                        if m.start() > pos: p2.add_run(para[pos:m.start()])
                        run_h = p2.add_run(para[m.start():m.end()]); run_h.bold=True; run_h.font.highlight_color = WD_COLOR_INDEX.RED
                        pos = m.end()
                    if pos < len(para): p2.add_run(para[pos:])
        browser.close()
    # save docs
    for kw, doc in docs.items():
        safe = kw.replace(' ', '_')
        out = f"{full_code}_{safe}_季度业绩.docx"
        doc.save(os.path.join(output_dir, out))
        logging.info(f"Saved combined DOCX for keyword {kw} [季度业绩]: {out}")

def download_original_reports(full_code: str, start_date: str = '2025-01-01', output_dir: str = 'original files'):
    """Download all PDFs from reports published after start_date."""
    os.makedirs(output_dir, exist_ok=True)
    page_index = 0  # 0-based page number for pagination
    while True:
        logging.info(f"Fetching page (original files): {page_index}")
        records = fetch_report_page(full_code, page_index, PAGE_SIZE)
        if not records:
            break
        for rec in records:
            pub_date = rec.get('publishDate', '').split()[0]
            logging.info(pub_date)
            if not pub_date or pub_date < start_date:
                continue
            raw_url = rec.get('url')
            logging.info(raw_url)
            if not raw_url:
                report_id = rec.get('reportId') or rec.get('id')
                detail_url = f"{DETAIL_BASE_URL}?id={report_id}&type={rec.get('type')}&storeId={STORE_ID}"
                logging.info(f"Downloading original HTML for record {rec.get('id')} from: {detail_url}")
                try:
                    resp = session.get(detail_url, timeout=10)
                    resp.raise_for_status()
                    content = resp.content
                except Exception as e:
                    logging.error(f"Error downloading HTML: {e}")
                    continue
                ext = '.html'
            else:
                # Detect HTML detail pages
                if '/report/detail' in raw_url:
                    html_url = raw_url
                    logging.info(f"Downloading original HTML for record {rec.get('id')} from: {html_url}")
                    try:
                        resp = session.get(html_url, timeout=10)
                        resp.raise_for_status()
                        content = resp.content
                    except Exception as e:
                        logging.error(f"Error downloading HTML: {e}")
                        continue
                    ext = '.html'
                else:
                    # PDF via preview wrapper
                    if "onlinePreview" in raw_url:
                        preview_url = raw_url
                    else:
                        b64 = base64.urlsafe_b64encode(raw_url.encode()).decode()
                        preview_url = (
                            f"https://file-view.comein.cn/onlinePreview?url={b64}"
                            "&officePreviewSwitchDisabled=true"
                            "&officePreviewType=pdf"
                            "&watermarkTxt="
                        )
                    logging.info(f"Downloading original PDF for record {rec.get('id')} from: {preview_url}")
                    content = download_pdf(preview_url)
                    ext = '.pdf'
            # save original file with code_title_author
            title = rec.get('title', '')
            safe_title = re.sub(r'\W+', '_', title).strip('_')
            author = rec.get('author', '')
            safe_author = re.sub(r'\W+', '_', author).strip('_')
            orig_fname = f"{full_code}_{safe_title}_{safe_author}{ext}"
            orig_path = os.path.join(output_dir, orig_fname)
            try:
                with open(orig_path, 'wb') as f:
                    f.write(content)
                logging.info(f"Saved original file {orig_fname}")
            except Exception as e:
                logging.error(f"Failed to save original file: {e}")
        if len(records) < PAGE_SIZE:
            break
        page_index += 1

if __name__ == '__main__':
    code = input('Enter company code (e.g. sz002738): ').strip() or 'sz002738'
    # input keywords interactively
    keywords = []
    while True:
        kw = input('Enter next keyword (blank to finish): ').strip()
        if not kw:
            break
        keywords.append(kw)
    if not keywords:
        print('No keywords provided. Exiting.')
        sys.exit(0)
    start_date = input('Enter earliest publish date (YYYY-MM-DD) or leave blank: ').strip() or None
    end_date = input('Enter latest publish date (YYYY-MM-DD) or leave blank: ').strip() or None
    crawl_company(code, keywords, start_date=start_date, end_date=end_date)
