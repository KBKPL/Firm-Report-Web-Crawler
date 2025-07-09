"""
Main module for PDF keyword search and exporting matching paragraphs to a docx file.
"""

import argparse
import io
import sys
import re
import requests
import urllib.parse
import base64
import subprocess
import fitz  # PyMuPDF for text extraction
import zipfile
from lxml import etree
from docx import Document
from bs4 import BeautifulSoup

def download_pdf(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        content_type = response.headers.get('content-type', '').lower()
        if 'html' in content_type:
            parsed = urllib.parse.urlparse(url)
            qs = urllib.parse.parse_qs(parsed.query)
            if 'url' in qs:
                b64 = qs['url'][0]
                real_url = base64.urlsafe_b64decode(b64).decode()
                response = requests.get(real_url)
                response.raise_for_status()
                return response.content
            else:
                print("Error: HTML response but no 'url' parameter for PDF.")
                sys.exit(1)
        return response.content
    except Exception as e:
        print(f"Error downloading PDF: {e}")
        sys.exit(1)

def split_chinese_sentences(text):
    # Merge lines and split only on Chinese punctuation (。！？)
    lines = text.split('\n')
    sentences = []
    buffer = ''
    for line in lines:
        line = line.strip()
        if not line:
            continue
        buffer += line
        if re.search(r'[。！？]$', buffer):
            sentences.append(buffer)
            buffer = ''
    if buffer:
        sentences.append(buffer)
    return sentences

def find_paragraphs_with_keyword(text, keyword):
    # Normalize newlines and split on one or more blank lines
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    blocks = [blk.strip() for blk in re.split(r'\n\s*\n+', text) if blk.strip()]
    keyword_lower = keyword.lower()
    # Return only blocks containing the keyword
    return [blk for blk in blocks if keyword_lower in blk.lower()]

def find_paragraphs_with_keyword_pymupdf(pdf_bytes, keyword):
    """
    Open PDF via PyMuPDF, extract text blocks, and filter by keyword.
    """
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        print(f"Error opening PDF via PyMuPDF: {e}")
        sys.exit(1)
    paras = []
    keyword_lower = keyword.lower()
    for page in doc:
        for b in page.get_text("blocks"):
            block_text = b[4].strip()
            if keyword_lower in block_text.lower():
                paras.append(block_text)
    return paras

def extract_text_with_pymupdf(pdf_bytes):
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        print(f"Error opening PDF via PyMuPDF: {e}")
        sys.exit(1)
    pages = [page.get_text("text") for page in doc]
    return "\n\n".join(pages)

def sanitize_text(text):
    # Remove control characters (except tab and newline)
    return ''.join(ch for ch in text if ch in ('\t','\n') or ord(ch) >= 32)

def find_paragraphs_in_docx_file(docx_path, keyword):
    # Read DOCX and group lines until a blank paragraph as one paragraph
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Error opening DOCX file {docx_path}: {e}")
        sys.exit(1)
    keyword_lower = keyword.lower()
    paragraphs = []
    buffer = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            buffer.append(txt)
        else:
            if buffer:
                full_para = ''.join(buffer)
                if keyword_lower in full_para.lower():
                    paragraphs.append(full_para)
                buffer = []
    if buffer:
        full_para = ''.join(buffer)
        if keyword_lower in full_para.lower():
            paragraphs.append(full_para)
    return paragraphs

def find_paragraphs_from_html(html_text, keyword):
    soup = BeautifulSoup(html_text, "lxml")
    keyword_lower = keyword.lower()
    paras = [p.get_text(strip=True) for p in soup.find_all('p') if p.get_text(strip=True)]
    return [p for p in paras if keyword_lower in p.lower()]

def write_paragraphs_to_docx(paragraphs, output_path, keyword):
    doc = Document()
    doc.add_heading(f'Paragraphs containing keyword: "{keyword}"', level=1)
    for p in paragraphs:
        clean_p = sanitize_text(p)
        if not clean_p.strip():
            continue  # skip empty after sanitization
        doc.add_paragraph(clean_p)
    try:
        doc.save(output_path)
    except Exception as e:
        print(f"Error writing DOCX file: {e}")
        sys.exit(1)

def main():
    parser = argparse.ArgumentParser(description='Search keyword in a PDF and export matching paragraphs to a docx file.')
    parser.add_argument('-u', '--url', help='URL of the PDF file')
    parser.add_argument('--docx-file', help='Path to pre-converted DOCX file to parse instead of PDF')
    parser.add_argument('-k', '--keyword', required=True, help='Keyword to search for')
    parser.add_argument('-o', '--output', default='output.docx', help='Output DOCX file name')
    args = parser.parse_args()
    # Either PDF URL or DOCX file must be provided
    if not args.url and not args.docx_file:
        parser.error('Provide either --url or --docx-file')

    keyword_lower = args.keyword.lower()
    if args.docx_file:
        print(f"Parsing DOCX file {args.docx_file} for '{args.keyword}'...")
        paragraphs = find_paragraphs_in_docx_file(args.docx_file, args.keyword)
    else:
        print(f"Fetching URL {args.url}...")
        try:
            resp = requests.get(args.url)
            resp.raise_for_status()
        except Exception as e:
            print(f"Error fetching URL: {e}")
            sys.exit(1)
        ctype = resp.headers.get("content-type", "").lower()
        # HTML page detected (not a PDF preview)
        qs = urllib.parse.parse_qs(urllib.parse.urlparse(args.url).query)
        if "html" in ctype and "url" not in qs:
            # Render page with Playwright to support JS
            try:
                from playwright.sync_api import sync_playwright
            except ImportError:
                print("Playwright not installed. Run 'pip install playwright' and 'playwright install'.")
                sys.exit(1)
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=True)
                page = browser.new_page()
                page.goto(args.url, timeout=60000)
                page.wait_for_load_state("networkidle")
                html_text = page.content()
                browser.close()
            # Extract text and search
            text = BeautifulSoup(html_text, "lxml").get_text("\n")
            if keyword_lower not in text.lower():
                print(f"No paragraphs found containing '{args.keyword}'.")
                sys.exit(0)
            paragraphs = find_paragraphs_with_keyword(text, args.keyword)
        else:
            # Handle PDF (direct or preview)
            if "html" in ctype and "url" in qs:
                b64 = qs.get("url")[0]
                real_url = base64.urlsafe_b64decode(b64).decode()
                resp = requests.get(real_url)
                resp.raise_for_status()
            pdf_bytes = resp.content
            # Pre-detect keyword in PDF
            if keyword_lower not in extract_text_with_pymupdf(pdf_bytes).lower():
                print(f"No paragraphs found containing '{args.keyword}'.")
                sys.exit(0)
            # Convert PDF to text via pdftext CLI
            sample_pdf = "sample.pdf"
            with open(sample_pdf, "wb") as f:
                f.write(pdf_bytes)
            print("Converting PDF to text via pdftext CLI...")
            try:
                subprocess.run(["pdftext", sample_pdf, "--out_path", "converted.txt"], check=True)
            except Exception as e:
                print(f"Error converting PDF to text: {e}")
                sys.exit(1)
            text = open("converted.txt", encoding="utf-8").read()
            paragraphs = find_paragraphs_with_keyword(text, args.keyword)

    if not paragraphs:
        print(f"No paragraphs found containing '{args.keyword}'.")
        sys.exit(0)

    print(f"Writing {len(paragraphs)} paragraphs to {args.output}...")
    write_paragraphs_to_docx(paragraphs, args.output, args.keyword)
    print("Done.")

if __name__ == '__main__':
    main()