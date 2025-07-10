"""
DOCX utilities: read/write paragraphs to/from DOCX.
"""
import sys
import logging
from docx import Document
from text_utils import sanitize_text
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.enum.text import WD_COLOR_INDEX

logger = logging.getLogger(__name__)

# Moved from src/crawler.py
def add_hyperlink(paragraph, url):
    """Insert a hyperlink into a Docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    i = OxmlElement('w:i')
    rPr.append(i)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = url
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_keyword_paragraphs(
    doc: Document,
    paragraphs: list[str],
    keyword: str,
    url: str,
    heading_fmt: str = "Location {idx}"
):
    """Append keyword-highlighted paragraphs with a hyperlink to a Document."""
    p = doc.add_paragraph()
    add_hyperlink(p, url)
    for idx, para in enumerate(paragraphs, start=1):
        doc.add_heading(heading_fmt.format(idx=idx), level=2)
        p2 = doc.add_paragraph()
        pattern = re.compile(re.escape(keyword), re.IGNORECASE)
        pos = 0
        for m in pattern.finditer(para):
            if m.start() > pos:
                p2.add_run(para[pos:m.start()])
            run_h = p2.add_run(para[m.start():m.end()])
            run_h.bold = True
            run_h.font.highlight_color = WD_COLOR_INDEX.RED
            pos = m.end()
        if pos < len(para):
            p2.add_run(para[pos:])
