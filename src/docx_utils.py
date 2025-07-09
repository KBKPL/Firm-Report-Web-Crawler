"""
DOCX utilities: read/write paragraphs to/from DOCX.
"""
import sys
import logging
from docx import Document
from text_utils import sanitize_text

logger = logging.getLogger(__name__)

def find_paragraphs_in_docx_file(docx_path: str, keyword: str) -> list[str]:
    """Extract paragraphs containing keyword from a .docx file."""
    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.error(f"Error opening DOCX file {docx_path}: {e}")
        sys.exit(1)
    keyword_lower = keyword.lower()
    paragraphs = []
    buffer = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            buffer.append(txt)
        else:
            if buffer:
                full_para = ''.join(buffer)
                if keyword_lower in full_para.lower():
                    paragraphs.append(full_para)
                buffer = []
    if buffer:
        full_para = ''.join(buffer)
        if keyword_lower in full_para.lower():
            paragraphs.append(full_para)
    return paragraphs

def write_paragraphs_to_docx(paragraphs: list[str], output_path: str, keyword: str) -> None:
    """Write paragraphs to a .docx file with headings."""
    doc = Document()
    count = len(paragraphs)
    doc.add_heading(f'Approximately {count} paragraph(s) including keyword: "{keyword}"', level=1)
    for idx, p in enumerate(paragraphs, start=1):
        doc.add_heading(f'Location {idx}', level=2)
        clean_p = sanitize_text(p)
        if not clean_p.strip():
            continue
        doc.add_paragraph(clean_p)
    try:
        doc.save(output_path)
    except Exception as e:
        logger.error(f"Error writing DOCX file: {e}")
        sys.exit(1)
